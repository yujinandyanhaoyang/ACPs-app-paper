from copy import deepcopy
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
import re
import zipfile
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
import subprocess
from docx.enum.style import WD_STYLE_TYPE


SRC = "/root/WORK/docs/thesis/本科毕业论文_合并稿.docx"
OUT = "/root/WORK/docs/thesis/本科毕业论文_合并稿_格式修正版.docx"
FIG_DIR = Path("/root/WORK/docs/thesis/generated_figures")

THESIS_TITLE_CN = "基于智能体协作的图书个性化推荐系统设计与实现"
THESIS_TITLE_EN = "Design and Implementation of Personalized Book Recommendation\nSystem Based onAgent Collaboration"


def set_run_font(run, eastasia="宋体", ascii_font="Times New Roman", size=None, bold=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), eastasia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def set_para_runs(paragraph, eastasia="宋体", ascii_font="Times New Roman", size=12, bold=None):
    for run in paragraph.runs:
        set_run_font(run, eastasia=eastasia, ascii_font=ascii_font, size=size, bold=bold)


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def delete_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    parent.remove(p)


def set_text(paragraph, text):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    return run


def add_page_field(paragraph):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(text)
    r.append(fld_end)


def add_toc_field(paragraph):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_end)


def apply_page_number_format(sectPr, fmt, start=None):
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sectPr.append(pg)
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))


def section_break_before(paragraph, source_sectPr, fmt, start=None):
    pPr = paragraph._p.get_or_add_pPr()
    sectPr = deepcopy(source_sectPr)
    apply_page_number_format(sectPr, fmt, start=start)
    pPr.append(sectPr)


def format_body_paragraph(paragraph, english=False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = paragraph.paragraph_format
    pf.first_line_indent = Cm(0.74)
    pf.left_indent = None
    pf.right_indent = None
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.line_spacing = 1.5
    if english:
        set_para_runs(paragraph, eastasia="Times New Roman", ascii_font="Times New Roman", size=12)
    else:
        set_para_runs(paragraph, eastasia="宋体", ascii_font="Times New Roman", size=12)


def format_heading1(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.first_line_indent = None
    pf.space_before = Pt(24)
    pf.space_after = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.line_spacing = 1.5
    set_para_runs(paragraph, eastasia="黑体", ascii_font="Times New Roman", size=16, bold=True)


def format_heading2(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.first_line_indent = None
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.line_spacing = 1.5
    set_para_runs(paragraph, eastasia="黑体", ascii_font="Times New Roman", size=14, bold=True)


def format_heading3(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.first_line_indent = None
    pf.space_before = Pt(6)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.line_spacing = 1.5
    set_para_runs(paragraph, eastasia="黑体", ascii_font="Times New Roman", size=12, bold=True)


def format_small_caption(paragraph, above=False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.first_line_indent = None
    pf.space_before = Pt(6 if above else 0)
    pf.space_after = Pt(6 if not above else 0)
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    set_para_runs(paragraph, eastasia="宋体", ascii_font="Times New Roman", size=9, bold=True)


def format_source_block(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.first_line_indent = None
    pf.left_indent = Cm(0.74)
    pf.right_indent = Cm(0.3)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0
    for run in paragraph.runs:
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(9)
        run.font.bold = False


def normalize_inline_list_text(text):
    if " - " not in text:
        return text
    if "：" in text:
        head, tail = text.split("：", 1)
        items = [p.strip() for p in tail.split(" - ") if p.strip()]
        if items:
            return head + "：\n" + "\n".join(f"- {p}" for p in items)
    items = [p.strip() for p in text.split(" - ") if p.strip()]
    if items:
        return "\n".join(f"- {p}" for p in items)
    return text


def normalize_algorithm_block(text):
    text = text.replace("输入:", "输入：").replace("输出:", "输出：")
    lines = [line.rstrip() for line in text.splitlines()]
    out = []
    in_io = False
    for line in lines:
        s = line.strip()
        if not s:
            out.append("")
            continue
        if s.startswith("输入：") or s.startswith("输出："):
            out.append(s)
            in_io = True
            continue
        if s.startswith("//"):
            in_io = False
            out.append(s)
            continue
        if in_io:
            out.append("  " + s)
        else:
            out.append(s)
    return "\n".join(out)


def generate_diagram_images():
    FIG_DIR.mkdir(exist_ok=True)
    def esc(s):
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    def svg_header(w, h):
        return [
            f'<svg xmlns="http://www.w3.org/2000/svg" width="{w}" height="{h}" viewBox="0 0 {w} {h}">',
            '<defs>',
            '<marker id="arrow" markerWidth="12" markerHeight="12" refX="10" refY="6" orient="auto" markerUnits="strokeWidth">',
            '<path d="M0,0 L12,6 L0,12 z" fill="#334155"/>',
            '</marker>',
            '<filter id="shadow" x="-20%" y="-20%" width="140%" height="140%">',
            '<feDropShadow dx="0" dy="2" stdDeviation="4" flood-color="#000000" flood-opacity="0.12"/>',
            '</filter>',
            '</defs>',
            '<rect width="100%" height="100%" fill="#ffffff"/>'
        ]

    def rect(x, y, w, h, fill="#ffffff", stroke="#334155", rx=18, sw=3):
        return f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="{rx}" ry="{rx}" fill="{fill}" stroke="{stroke}" stroke-width="{sw}" filter="url(#shadow)"/>'

    def text(x, y, lines, size=28, weight="600", fill="#0f172a"):
        if isinstance(lines, str):
            lines = [lines]
        parts = [f'<text x="{x}" y="{y}" font-family="DejaVu Sans, Arial, sans-serif" font-size="{size}" font-weight="{weight}" text-anchor="middle" fill="{fill}">']
        for i, line in enumerate(lines):
            dy = "0" if i == 0 else "1.35em"
            parts.append(f'<tspan x="{x}" dy="{dy}">{esc(line)}</tspan>')
        parts.append('</text>')
        return "".join(parts)

    def line(x1, y1, x2, y2, dashed=False):
        dash = ' stroke-dasharray="10 8"' if dashed else ""
        return f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="#334155" stroke-width="4"{dash} marker-end="url(#arrow)"/>'

    # fig3-1 as SVG
    w, h = 2200, 1300
    parts = svg_header(w, h)
    parts += [
        rect(860, 90, 480, 120, fill="#E0F2FE"),
        text(1100, 140, ["Reading Concierge", "Leader"], 34, "700"),
        rect(80, 310, 360, 120, fill="#F8FAFC"),
        text(260, 360, ["User Request"], 30, "700"),
        rect(250, 550, 340, 130, fill="#EEF2FF"),
        text(420, 605, ["reader_profile_agent", "Profile Proposal"], 28),
        rect(690, 550, 340, 130, fill="#F5F3FF"),
        text(860, 605, ["book_content_agent", "Content Proposal"], 28),
        rect(1130, 550, 420, 130, fill="#FEF3C7"),
        text(1340, 605, ["recommendation_decision_agent", "Strategy Arbitration"], 28),
        rect(1650, 550, 380, 130, fill="#DCFCE7"),
        text(1840, 605, ["recommendation_engine_agent", "Recommendation Execution"], 28),
        rect(860, 820, 480, 130, fill="#FCE7F3"),
        text(1100, 875, ["feedback_agent", "Feedback Update"], 28),
        rect(180, 1050, 780, 130, fill="#F8FAFC"),
        text(570, 1105, ["ACPs Protocol Layer", "AIC / ACS / ATR / AIA / ADP / DSP / AIP"], 28),
        rect(1190, 1050, 830, 130, fill="#F8FAFC"),
        text(1605, 1105, ["Data & Model Layer", "Behavior Store / FAISS / CF / Arm Records"], 28),
        line(440, 370, 860, 150),
        line(1100, 210, 420, 550),
        line(1100, 210, 860, 550),
        line(1100, 210, 1340, 550),
        line(1550, 615, 1650, 615),
        line(1840, 680, 1100, 820),
        line(1100, 950, 420, 680),
        line(1100, 950, 860, 680),
        line(1100, 950, 1340, 680),
        line(1100, 950, 1840, 680),
    ]
    for x in [1100, 420, 860, 1340, 1840, 1100]:
        y1, y2 = (210, 1050) if x != 1840 else (680, 1050)
        parts.append(line(x, y1, x, y2, dashed=True))
    parts.append('</svg>')
    (FIG_DIR / "fig3_1.svg").write_text("\n".join(parts), encoding="utf-8")
    subprocess.run(["rsvg-convert", str(FIG_DIR / "fig3_1.svg"), "-o", str(FIG_DIR / "fig3_1.png"), "-w", "2600"], check=True)

    # fig3-2 as SVG
    w, h = 2400, 1500
    xs = [180, 520, 860, 1200, 1540, 1880, 2220]
    labels = ["User", "Leader", "Profile Agent", "Content Agent", "Arbiter", "Engine", "Feedback"]
    parts = svg_header(w, h)
    for x, label in zip(xs, labels):
        parts.append(text(x, 90, [label], 30, "700"))
        parts.append(f'<line x1="{x}" y1="140" x2="{x}" y2="1360" stroke="#94A3B8" stroke-width="3" stroke-dasharray="10 8"/>')
    steps = [
        (0, 1, 190, "Submit Query"),
        (1, 2, 310, "Request Profile Proposal"),
        (1, 3, 390, "Request Content Proposal"),
        (2, 1, 500, "Return Profile Vector / Confidence"),
        (3, 1, 580, "Return Content Proposal / Divergence"),
        (1, 4, 700, "Submit Arbitration Input"),
        (4, 1, 810, "Return Strategy Parameters"),
        (1, 5, 930, "Submit Execution Parameters"),
        (5, 1, 1040, "Return Recommendation List / Explanation"),
        (0, 6, 1180, "View / Click / Rate / Finish"),
        (6, 4, 1280, "Reward Update"),
        (6, 2, 1340, "Trigger Profile Refresh"),
        (6, 5, 1400, "Trigger Retraining"),
    ]
    for a, b, y, label in steps:
        x1, x2 = xs[a], xs[b]
        parts.append(line(x1, y, x2, y))
        parts.append(text((x1 + x2) / 2, y - 18, [label], 22, "500"))
    parts.append('</svg>')
    (FIG_DIR / "fig3_2.svg").write_text("\n".join(parts), encoding="utf-8")
    subprocess.run(["rsvg-convert", str(FIG_DIR / "fig3_2.svg"), "-o", str(FIG_DIR / "fig3_2.png"), "-w", "2800"], check=True)

    title_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 28)
    label_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 18)
    small_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 15)

    palette = ["#355C7D", "#6C8EAD", "#C06C84", "#F67280", "#99B898", "#F8B195"]

    def nice_chart_canvas(w=1500, h=900, title=""):
        img = Image.new("RGB", (w, h), "white")
        draw = ImageDraw.Draw(img)
        draw.rounded_rectangle((20, 20, w - 20, h - 20), 24, outline="#D0D7DE", width=2, fill="white")
        draw.text((w / 2, 55), title, fill="#111827", font=title_font, anchor="mm")
        return img, draw

    def grouped_bars(draw, origin, size, categories, series_names, data, colors, y_label="", y_max=None):
        x0, y0 = origin
        w, h = size
        y_max = y_max or max(max(s) for s in data) * 1.15
        plot_top = y0 + 20
        plot_bottom = y0 + h - 70
        plot_left = x0 + 70
        plot_right = x0 + w - 20
        draw.line((plot_left, plot_top, plot_left, plot_bottom), fill="#111827", width=2)
        draw.line((plot_left, plot_bottom, plot_right, plot_bottom), fill="#111827", width=2)
        for k in range(6):
            val = y_max * k / 5
            y = plot_bottom - (plot_bottom - plot_top) * k / 5
            draw.line((plot_left, y, plot_right, y), fill="#E5E7EB", width=1)
            txt = f"{val:.4f}" if y_max < 1 else f"{val:.2f}"
            draw.text((plot_left - 10, y), txt, fill="#4B5563", font=small_font, anchor="rm")
        group_w = (plot_right - plot_left) / len(categories)
        bar_w = group_w / (len(series_names) + 1.2)
        for i, cat in enumerate(categories):
            cx = plot_left + group_w * (i + 0.5)
            draw.text((cx, plot_bottom + 30), cat, fill="#111827", font=small_font, anchor="ma")
            for j, series in enumerate(data):
                val = series[i]
                left = plot_left + i * group_w + 0.15 * group_w + j * bar_w
                right = left + bar_w * 0.9
                top = plot_bottom - (val / y_max) * (plot_bottom - plot_top)
                draw.rounded_rectangle((left, top, right, plot_bottom), 6, fill=colors[j], outline=colors[j])
                draw.text(((left + right) / 2, top - 8), f"{val:.4f}" if val < 1 else f"{val:.2f}", fill="#374151", font=small_font, anchor="ms")
        lx = plot_right - 260
        ly = plot_top + 5
        for j, name in enumerate(series_names):
            draw.rectangle((lx, ly + j * 28, lx + 20, ly + 16 + j * 28), fill=colors[j], outline=colors[j])
            draw.text((lx + 28, ly + 8 + j * 28), name, fill="#111827", font=label_font, anchor="lm")

    def simple_bars(draw, origin, size, categories, values, color, title=None, y_max=None):
        x0, y0 = origin
        w, h = size
        if title:
            draw.text((x0 + w / 2, y0 + 10), title, fill="#111827", font=label_font, anchor="ma")
        plot_top = y0 + 40
        plot_bottom = y0 + h - 55
        plot_left = x0 + 60
        plot_right = x0 + w - 20
        y_max = y_max or max(values) * 1.12
        draw.line((plot_left, plot_top, plot_left, plot_bottom), fill="#111827", width=2)
        draw.line((plot_left, plot_bottom, plot_right, plot_bottom), fill="#111827", width=2)
        for k in range(6):
            val = y_max * k / 5
            y = plot_bottom - (plot_bottom - plot_top) * k / 5
            draw.line((plot_left, y, plot_right, y), fill="#E5E7EB", width=1)
            txt = f"{val:.3f}" if y_max < 2 else f"{val:.1f}"
            draw.text((plot_left - 8, y), txt, fill="#4B5563", font=small_font, anchor="rm")
        group_w = (plot_right - plot_left) / len(categories)
        bar_w = group_w * 0.58
        for i, (cat, val) in enumerate(zip(categories, values)):
            left = plot_left + i * group_w + (group_w - bar_w) / 2
            right = left + bar_w
            top = plot_bottom - (val / y_max) * (plot_bottom - plot_top)
            draw.rounded_rectangle((left, top, right, plot_bottom), 6, fill=color, outline=color)
            draw.text(((left + right) / 2, top - 8), f"{val:.3f}" if val < 2 else f"{val:.2f}", fill="#374151", font=small_font, anchor="ms")
            draw.text(((left + right) / 2, plot_bottom + 22), cat, fill="#111827", font=small_font, anchor="ma")

    # Figure 4-1
    img, draw = nice_chart_canvas(title="Accuracy Metrics Comparison")
    categories = ["PopRec", "Content", "Hybrid", "LightGCN", "Reading Concierge"]
    grouped_bars(
        draw,
        (90, 120),
        (1320, 680),
        categories,
        ["Precision@10", "Recall@10", "NDCG@10"],
        [
            [0.0006, 0.00105, 0.0013, 0.0091, 0.0100],
            [0.00179, 0.002725, 0.003578, 0.041394, 0.0750],
            [0.001532, 0.002406, 0.003208, 0.023993, 0.037785],
        ],
        [palette[0], palette[1], palette[3]],
    )
    img.save(FIG_DIR / "fig4_1.png")

    # Figure 4-2
    img, draw = nice_chart_canvas(title="Ablation Study Accuracy Comparison")
    categories = ["Full", "w/o CF", "w/o Content", "w/o MMR", "w/o Arbiter"]
    grouped_bars(
        draw,
        (90, 120),
        (1320, 680),
        categories,
        ["Precision@10", "Recall@10", "NDCG@10"],
        [
            [0.001000, 0.000200, 0.001000, 0.000600, 0.001000],
            [0.002958, 0.000125, 0.003044, 0.001458, 0.002833],
            [0.002309, 0.000440, 0.002381, 0.001661, 0.001639],
        ],
        [palette[0], palette[1], palette[3]],
    )
    img.save(FIG_DIR / "fig4_2.png")

    # Figure 4-3
    img, draw = nice_chart_canvas(title="Ablation Study Diversity Comparison")
    cats = ["Full", "w/o CF", "w/o Content", "w/o MMR", "w/o Arbiter"]
    simple_bars(draw, (60, 130), (430, 600), cats, [0.969326, 0.954363, 0.968104, 0.810889, 0.978341], palette[0], "ILD", y_max=1.0)
    simple_bars(draw, (520, 130), (430, 600), cats, [0.000670, 0.000879, 0.000291, 0.000751, 0.000372], palette[1], "Coverage", y_max=0.0010)
    simple_bars(draw, (980, 130), (430, 600), cats, [11.223756, 13.078116, 10.566534, 11.401945, 10.732025], palette[3], "Novelty", y_max=14.0)
    img.save(FIG_DIR / "fig4_3.png")

    # Figure 4-4
    img, draw = nice_chart_canvas(title="Performance across User Activity Levels")
    categories = ["Low\n(5-10)", "Medium\n(11-20)", "High\n(20+)"]
    grouped_bars(
        draw,
        (90, 120),
        (1320, 680),
        categories,
        ["Precision@10", "Recall@10", "NDCG@10"],
        [
            [0.0009, 0.0011, 0.0012],
            [0.002125, 0.003188, 0.003562],
            [0.001978, 0.002456, 0.002593],
        ],
        [palette[0], palette[1], palette[3]],
    )
    img.save(FIG_DIR / "fig4_4.png")


def format_reference(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = paragraph.paragraph_format
    pf.first_line_indent = Cm(-0.74)
    pf.left_indent = Cm(0.74)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    set_para_runs(paragraph, eastasia="宋体", ascii_font="Times New Roman", size=9)


def format_table(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pf = p.paragraph_format
                pf.first_line_indent = None
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = 1.5
                pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                set_para_runs(p, eastasia="宋体", ascii_font="Times New Roman", size=9)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    spec = {
        "top": ("12", "single"),
        "left": ("0", "nil"),
        "bottom": ("12", "single"),
        "right": ("0", "nil"),
        "insideH": ("6", "single"),
        "insideV": ("0", "nil"),
    }
    for edge, (sz, val) in spec.items():
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), val)
        element.set(qn("w:sz"), sz)
        element.set(qn("w:color"), "000000")


def main():
    doc = Document(SRC)
    generate_diagram_images()

    for p in doc.paragraphs:
        if "基于ACPs的图书个性化推荐系统设计与实现" in p.text:
            p.text = p.text.replace("基于ACPs的图书个性化推荐系统设计与实现", THESIS_TITLE_CN)

    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.5)

    header = sec.header
    hp = header.paragraphs[0]
    set_text(hp, "北京邮电大学本科毕业设计（论文）")
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(0)
    set_para_runs(hp, eastasia="宋体", ascii_font="Times New Roman", size=9)

    footer = sec.footer
    fp = footer.paragraphs[0]
    clear_paragraph(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)
    set_para_runs(fp, eastasia="宋体", ascii_font="Times New Roman", size=9)

    # Remove residual pseudo-cover lines from the generated file.
    delete_paragraph(doc.paragraphs[2])
    delete_paragraph(doc.paragraphs[1])
    title_para = doc.paragraphs[0]
    abs_heading_cn = doc.paragraphs[1]
    first_abs_body_text = doc.paragraphs[2].text

    first = title_para
    p = first.insert_paragraph_before("北 京 邮 电 大 学")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    set_para_runs(p, eastasia="黑体", ascii_font="Times New Roman", size=22, bold=True)
    p = first.insert_paragraph_before("本 科 毕 业 设 计（论 文）")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)
    set_para_runs(p, eastasia="黑体", ascii_font="Times New Roman", size=26, bold=True)
    first.insert_paragraph_before("")
    first.insert_paragraph_before("")
    p = first.insert_paragraph_before(THESIS_TITLE_CN)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    set_para_runs(p, eastasia="黑体", ascii_font="Times New Roman", size=22, bold=True)
    first.insert_paragraph_before("")
    cover_lines = [
        "姓    名        XXX",
        "学    院        XXX学院",
        "专    业        XXX专业",
        "班    级        2023XXXX",
        "学    号        2023XXXXXX",
        "指导教师      XXX",
        "202X 年 6 月",
    ]
    for text in cover_lines:
        p = first.insert_paragraph_before(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_runs(p, eastasia="宋体", ascii_font="Times New Roman", size=16)
    p = first.insert_paragraph_before("")
    p.add_run().add_break(WD_BREAK.PAGE)
    set_para_runs(p)

    # Statement and authorization pages.
    first = title_para
    p = first.insert_paragraph_before("本科毕业设计（论文）诚信声明")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_runs(p, eastasia="黑体", ascii_font="Times New Roman", size=22, bold=True)
    body = [
        "本人声明所呈交的毕业设计（论文）是本人在指导教师指导下独立完成的研究成果。除文中已经明确标注引用的内容外，论文不包含任何他人已经发表或撰写过的研究成果，也不包含本人为获得其他学位或证书而使用过的材料。",
        "对本文的研究做出重要贡献的个人和集体，均已在文中以明确方式标明。本人完全意识到本声明的法律后果由本人承担。",
        "本人签名：                         日期：XXXX年XX月XX日",
    ]
    for text in body:
        p = first.insert_paragraph_before(text)
        format_body_paragraph(p)
    p = first.insert_paragraph_before("")
    p.add_run().add_break(WD_BREAK.PAGE)
    set_para_runs(p)

    first = title_para
    p = first.insert_paragraph_before("关于论文使用授权的说明")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_runs(p, eastasia="黑体", ascii_font="Times New Roman", size=22, bold=True)
    body = [
        "本人完全了解并同意北京邮电大学有关保留、使用学位论文的规定。学校有权保留并向有关部门或机构送交论文的复印件和电子版，允许论文被查阅和借阅，可以采用影印、缩印或扫描等复制手段保存、汇编本论文。",
        "本人授权北京邮电大学将本论文的全部或部分内容用于教学、科研和学术交流，但涉及保密内容的部分除外。",
        "本人签名：                         日期：XXXX年XX月XX日",
        "导师签名：                         日期：XXXX年XX月XX日",
    ]
    for text in body:
        p = first.insert_paragraph_before(text)
        format_body_paragraph(p)
    p = first.insert_paragraph_before("")
    p.add_run().add_break(WD_BREAK.PAGE)
    set_para_runs(p)

    # Chinese abstract page formatting.
    title_para.text = THESIS_TITLE_CN
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_runs(title_para, eastasia="黑体", ascii_font="Times New Roman", size=22, bold=True)
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(12)

    abs_heading_cn.text = "摘要"
    abs_heading_cn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_runs(abs_heading_cn, eastasia="黑体", ascii_font="Times New Roman", size=16, bold=True)
    abs_heading_cn.paragraph_format.space_before = Pt(12)
    abs_heading_cn.paragraph_format.space_after = Pt(12)

    start_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text == first_abs_body_text)
    for idx in range(start_idx, min(start_idx + 5, len(doc.paragraphs))):
        format_body_paragraph(doc.paragraphs[idx])

    for p in doc.paragraphs:
        if p.text.strip().startswith("关键词"):
            p.text = "关键词  个性化推荐 ACPs协议 多智能体协作 图书推荐 上下文赌博机"
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = None
            set_para_runs(p, eastasia="宋体", ascii_font="Times New Roman", size=12, bold=True)
            break

    # English abstract title before ABSTRACT.
    abs_para = next(p for p in doc.paragraphs if p.text.strip() == "ABSTRACT")
    p = abs_para.insert_paragraph_before("")
    p.add_run().add_break(WD_BREAK.PAGE)
    p = abs_para.insert_paragraph_before(THESIS_TITLE_EN)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_runs(p, eastasia="Times New Roman", ascii_font="Times New Roman", size=22, bold=True)
    abs_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_runs(abs_para, eastasia="Times New Roman", ascii_font="Times New Roman", size=16, bold=True)
    abs_para.paragraph_format.space_before = Pt(12)
    abs_para.paragraph_format.space_after = Pt(12)

    english = False
    for p in doc.paragraphs:
        if p.text.strip() == "ABSTRACT":
            english = True
            continue
        if p.text.strip().startswith("第一章 "):
            english = False
        if english:
            if p.text.strip().startswith("KEY WORDS"):
                p.text = "KEY WORDS  personalized recommendation ACPs protocol multi-agent collaboration book recommendation contextual bandit"
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.first_line_indent = None
                set_para_runs(p, eastasia="Times New Roman", ascii_font="Times New Roman", size=12, bold=True)
            elif p.text.strip():
                format_body_paragraph(p, english=True)

    # Insert TOC page and section break before chapter 1.
    chapter1 = next(p for p in doc.paragraphs if p.text.strip() == "第一章 绪论")
    p = chapter1.insert_paragraph_before("目录")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_runs(p, eastasia="黑体", ascii_font="Times New Roman", size=16, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p = chapter1.insert_paragraph_before("")
    add_toc_field(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = None
    set_para_runs(p, eastasia="宋体", ascii_font="Times New Roman", size=12)
    p = chapter1.insert_paragraph_before("")
    p.add_run().add_break(WD_BREAK.PAGE)

    section_break_before(chapter1, doc.sections[0]._sectPr, fmt="upperRoman", start=1)
    apply_page_number_format(doc.sections[0]._sectPr, "decimal", start=1)

    in_reference = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if re.match(r"^第[一二三四五六七八九十]+章 ", text) or text in {"参考文献", "附录A 缩略语表", "附录B 核心算法伪代码", "附录C 关键API接口说明", "致  谢", "攻读学位期间取得的创新成果"}:
            format_heading1(p)
            in_reference = text == "参考文献"
            continue
        if re.match(r"^\d+\.\d+\s", text) or re.match(r"^[A-Z]\.\d+\s", text) or re.match(r"^算法[A-Z]-\d+\s", text):
            if text.startswith("算法"):
                format_small_caption(p, above=True)
            else:
                format_heading2(p)
            in_reference = False
            continue
        if re.match(r"^\d+\.\d+\.\d+\s", text) or re.match(r"^[A-Z]\.\d+\.\d+\s", text):
            format_heading3(p)
            in_reference = False
            continue
        if re.match(r"^\[\d+\]", text) and in_reference:
            format_reference(p)
            continue
        if re.match(r"^图\d+-\d+\s", text):
            format_small_caption(p, above=False)
            continue
        if re.match(r"^表\d+-\d+\s", text) or re.match(r"^续表\d+-\d+\s", text):
            format_small_caption(p, above=True)
            continue
        if text.startswith("关键词"):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = None
            set_para_runs(p, eastasia="宋体", ascii_font="Times New Roman", size=12, bold=True)
            continue
        if text.startswith("KEY WORDS"):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = None
            set_para_runs(p, eastasia="Times New Roman", ascii_font="Times New Roman", size=12, bold=True)
            continue
        if in_reference:
            format_reference(p)
        else:
            format_body_paragraph(p, english=False)

    duplicate_figure_lines = {
        "图 4-1 各方法准确性指标对比",
        "图 4-2 消融实验准确性对比",
        "图 4-3 消融实验多样性对比",
        "图 4-4 不同用户活跃度性能对比",
    }
    duplicate_indices = [i for i, p in enumerate(doc.paragraphs) if p.text.strip() in duplicate_figure_lines]
    for i in reversed(duplicate_indices):
        delete_paragraph(doc.paragraphs[i])

    for p in doc.paragraphs:
        text = p.text.strip()
        if re.match(r"^图\d+-\d+\s", text):
            p.text = re.sub(r"^图(\d+-\d+)\s*", r"图 \1 ", text)
            text = p.text.strip()
        elif re.match(r"^表\s*\d+-\d+\s", text):
            p.text = re.sub(r"^表\s*(\d+-\d+)\s*", r"表 \1 ", text)
            text = p.text.strip()
        elif re.match(r"^算法\s*\d+-\d+\s", text):
            p.text = re.sub(r"^算法\s*(\d+-\d+)\s*", r"算法 \1 ", text)
            text = p.text.strip()
        elif re.match(r"^算法B-\d+\s", text):
            p.text = re.sub(r"^算法(B-\d+)\s*", r"算法 \1 ", text)
            text = p.text.strip()
        if text == "目录":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_para_runs(p, eastasia="黑体", ascii_font="Times New Roman", size=16, bold=True)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
        elif re.match(r"^图\s*\d+-\d+\s", text):
            format_small_caption(p, above=False)
        elif re.match(r"^表\s*\d+-\d+\s", text):
            format_small_caption(p, above=True)
        elif re.match(r"^算法\s*\d+-\d+\s", text) or re.match(r"^算法[A-Z]-\d+\s", text):
            format_small_caption(p, above=True)
        elif text.startswith("[") and re.match(r"^\[\d+\]", text):
            format_reference(p)
        elif p.style.name == "Source Code":
            if text.startswith("输入:") or text.startswith("输入："):
                p.text = normalize_algorithm_block(text)
            format_source_block(p)

    # Appendix algorithm titles should not look like normal section headings.
    for p in doc.paragraphs:
        text = p.text.strip()
        if re.match(r"^算法B-\d+\s", text):
            format_small_caption(p, above=True)
        elif re.match(r"^算法\s*3-\d+\s", text):
            format_small_caption(p, above=True)

    # Improve appendix labels and bullet-like lines that came from markdown.
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("输入：") or text.startswith("输出：") or text.startswith("说明：") or text.startswith("功能：") or text.startswith("输入参数：") or text.startswith("输出参数：") or text.startswith("示例调用：") or text.startswith("反馈事件格式：") or text.startswith("典型的推荐请求调用流程如下："):
            p.text = normalize_inline_list_text(text)
            format_body_paragraph(p)
            p.paragraph_format.first_line_indent = None
        elif text.startswith("常见错误码："):
            p.text = normalize_inline_list_text(text)
            format_body_paragraph(p)
            p.paragraph_format.first_line_indent = None

    for p in doc.paragraphs:
        text = p.text.strip()
        if text in {"致  谢", "攻读学位期间取得的创新成果"}:
            format_heading1(p)
        elif text in {"论文", "专利", "竞赛"}:
            format_heading3(p)

    for table in doc.tables:
        format_table(table)
        set_table_borders(table)

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith("flowchart LR"):
            clear_paragraph(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(FIG_DIR / "fig3_1.png"), width=Cm(15.5))
        elif text.startswith("sequenceDiagram"):
            clear_paragraph(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(FIG_DIR / "fig3_2.png"), width=Cm(16))

    figure_map = {
        "图4-1 准确性指标对比": ("fig4_1.png", 15.8),
        "图4-2 消融实验准确性对比": ("fig4_2.png", 15.8),
        "图4-3 消融实验多样性对比": ("fig4_3.png", 16.2),
        "图4-4 用户活跃度性能对比": ("fig4_4.png", 15.8),
    }
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text in figure_map and i > 0:
            img_name, width_cm = figure_map[text]
            prev = doc.paragraphs[i - 1]
            clear_paragraph(prev)
            prev.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = prev.add_run()
            run.add_picture(str(FIG_DIR / img_name), width=Cm(width_cm))

    # Add closing sections at the real end of the document.
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    p = doc.add_paragraph("致  谢")
    format_heading1(p)
    p = doc.add_paragraph("感谢指导教师、课题组同学、家人与朋友在论文撰写与系统实现过程中给予的帮助与支持。本文相关研究工作得以顺利推进，离不开各方在研究思路、实验环境和论文修改方面的持续支持。")
    format_body_paragraph(p)
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    p = doc.add_paragraph("攻读学位期间取得的创新成果")
    format_heading1(p)
    items = [
        "论文",
        "[1] 暂无。",
        "专利",
        "[1] 暂无。",
        "竞赛",
        "[1] 暂无。",
    ]
    for text in items:
        p = doc.add_paragraph(text)
        if text in {"论文", "专利", "竞赛"}:
            format_heading3(p)
        else:
            format_body_paragraph(p)

    doc.save(OUT)

    # Fix section page-number formats directly in document.xml.
    with zipfile.ZipFile(OUT, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}
    xml = files["word/document.xml"].decode("utf-8")
    sects = re.findall(r"<w:sectPr[\s\S]*?</w:sectPr>", xml)
    if len(sects) >= 2:
        first = re.sub(
            r"<w:pgNumType[^>]*/>",
            '<w:pgNumType w:start="1" w:fmt="upperRoman"/>',
            sects[0],
            count=1,
        )
        second = re.sub(
            r"<w:pgNumType[^>]*/>",
            '<w:pgNumType w:start="1" w:fmt="decimal"/>',
            sects[1],
            count=1,
        )
        xml = xml.replace(sects[0], first, 1)
        xml = xml.replace(sects[1], second, 1)
    files["word/document.xml"] = xml.encode("utf-8")
    with zipfile.ZipFile(OUT, "w") as zout:
        for name, data in files.items():
            zout.writestr(name, data)


if __name__ == "__main__":
    main()
